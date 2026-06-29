# -*- coding: utf-8 -*-
import json, pandas as pd, glob
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FONT='Times New Roman'; BLACK=RGBColor(0,0,0)
def fmt(x,dec=3):
    try: return f'{float(x):.{dec}f}'.replace('.',',')
    except: return str(x)

dm=pd.read_csv('experiments/results/metrics_reports/descriptive_means.csv').set_index('method')
rk=pd.read_csv('experiments/results/metrics_reports/ranking_methods.csv').set_index('Unnamed: 0')
fr=pd.read_csv('experiments/results/metrics_reports/friedman_results.csv')
ds=pd.read_csv('experiments/results/metrics_reports/detection_summary.csv').set_index('method')
abl=json.load(open('experiments/results/metrics_reports/ablation_multiescala.json'))

ORDER=[('Promedio','Promedio'),('PiramideLaplace','Pirámide de Laplace'),('Curvelet','Curvelet'),
('TopHat_disk_L3','Torre TH disco L3'),('TopHat_square_L3','Torre TH cuadrado L3'),
('TopHat_cross_L3','Torre TH cruz L3'),('TopHat_disk_L5','Torre TH disco L5 (anterior)'),
('TopHat_disk_L3_BTH','Torre TH disco L3 +BTH'),('TopHat_square_L3_BTH','Torre TH cuadrado L3 +BTH'),
('TopHat_cross_L3_BTH','Torre TH cruz L3 +BTH'),('TopHat_disk_L5_BTH','Torre TH disco L5 +BTH'),
('TopHat_Optimo','Óptimo monoescala'),('Optimo_Multiescala','Óptimo multiescala (propuesta)')]

doc=Document()
sec=doc.sections[0]
for m in ('top_margin','bottom_margin','left_margin','right_margin'): setattr(sec,m,Inches(1.0))
nm=doc.styles['Normal']; nm.font.name=FONT; nm.font.size=Pt(12)
nm.paragraph_format.space_after=Pt(6); nm.paragraph_format.line_spacing=1.15

def para(txt,size=12,bold=False,italic=False,align=WD_ALIGN_PARAGRAPH.LEFT,after=6):
    p=doc.add_paragraph(); p.alignment=align; p.paragraph_format.space_after=Pt(after)
    r=p.add_run(txt); r.font.name=FONT; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=BLACK
    return p
def heading(txt,size=13,before=12):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(txt); r.font.name=FONT; r.font.size=Pt(size); r.font.bold=True; r.font.color.rgb=BLACK
    return p

def table(headers, rows):
    t=doc.add_table(rows=1+len(rows), cols=len(headers))
    bd=OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        b=OxmlElement('w:'+e); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:color'),'000000'); bd.append(b)
    t._tbl.tblPr.append(bd)
    for ci,h in enumerate(headers):
        c=t.rows[0].cells[ci]; c.paragraphs[0].text=''
        r=c.paragraphs[0].add_run(h); r.font.name=FONT; r.font.size=Pt(10); r.font.bold=True; r.font.color.rgb=BLACK
    for ri,row in enumerate(rows,1):
        for ci,val in enumerate(row):
            c=t.rows[ri].cells[ci]; c.paragraphs[0].text=''
            r=c.paragraphs[0].add_run(str(val)); r.font.name=FONT; r.font.size=Pt(10); r.font.color.rgb=BLACK
            if ci>0: c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    return t

# ---- título ----
para('Estudio de resultados de los experimentos de fusión VIS/IR', size=15, bold=True, after=2)
para('Tablas cuantitativas (12 métricas) y comparación cualitativa por escena', size=12, italic=True, after=10)
para('TFM — Maestría en Ciencias de Datos, UCOM. Dataset: TNO Image Fusion (20 pares VIS/IR). '
 'Se evalúan 13 métodos: tres baselines (Promedio, Pirámide de Laplace, Curvelet), ocho configuraciones de la '
 'Torre Top-Hat (método anterior, en modo WTH y WTH+BTH), el óptimo monoescala y el óptimo multiescala (propuesta, '
 'con PSO: n=6, r≈2,89, m=0,10). Las métricas son sin referencia: seis clásicas (EN, SD, FE, MG, MI_vis, MI_ir) y '
 'seis estándar de calidad (SF, Qabf, Nabf, SSIM, SCD, VIF).')

heading('1. Resultados cuantitativos', 14)

heading('1.1 Métricas clásicas de actividad e información (medias, n = 20)')
table(['Método','EN','SD','FE','MG','MI_vis','MI_ir'],
      [[name]+[fmt(dm.loc[k,m]) for m in ['EN','SD','FE','MG','MI_vis','MI_ir']] for k,name in ORDER])
para('En todas, mayor es mejor. El disco L5 +BTH maximiza EN, SD y MG (contraste y actividad), pero a costa de la calidad (ver 1.2).', size=10, after=10)

heading('1.2 Métricas estándar de calidad de fusión (medias, n = 20)')
table(['Método','SF','Qabf','Nabf','SSIM','SCD','VIF'],
      [[name]+[fmt(dm.loc[k,m]) for m in ['SF','Qabf','Nabf','SSIM','SCD','VIF']] for k,name in ORDER])
para('Aquí mayor es mejor salvo Nabf (menor es mejor). El óptimo multiescala obtiene el mayor Qabf (0,514) y SCD (1,453), '
 'el menor Nabf entre los métodos no triviales (0,065) y el segundo SSIM (0,775).', size=10, after=10)

heading('1.3 Ranking por rangos promedio (1 = mejor)')
rows=[]
for k,name in ORDER:
    if k in rk.index:
        rows.append([name]+[fmt(rk.loc[k,m],1) for m in ['Qabf','Nabf','SSIM','SCD','VIF','SF']]+[fmt(rk.loc[k,'avg_rank'],2)])
rows.sort(key=lambda r: float(r[-1].replace(',','.')))
table(['Método','Qabf','Nabf','SSIM','SCD','VIF','SF','Rango global'],rows)
para('Rango promedio sobre las doce métricas (menor es mejor). La Pirámide de Laplace encabeza el agregado por su ventaja en '
 'contraste; el óptimo multiescala domina las métricas de calidad.', size=10, after=10)

heading('1.4 Significancia estadística — prueba de Friedman')
table(['Métrica','χ² de Friedman','p-valor','Significativo (α=0,05)'],
      [[r['metric'],fmt(r['chi2'],2),f"{r['p_value']:.2e}".replace('.',','),'Sí' if r['significant_05'] else 'No'] for _,r in fr.iterrows()])
para('Todas las métricas son altamente significativas (p mucho menor que 0,001): las diferencias entre métodos no se deben al azar.', size=10, after=10)

heading('1.5 Ablación del método óptimo multiescala (5 escenas representativas)')
amap=[('Monoescala (n=1)','Monoescala (n = 1)'),('n=2','Cascada n = 2'),('n=4','Cascada n = 4'),
 ('Disco + 4 lineas — propuesta (n=6)','Cascada n = 6 (propuesta)'),('n=8','Cascada n = 8'),('Disco solo (n=6)','n = 6, solo disco')]
table(['Variante','Qabf','SSIM','SCD','Nabf'],
      [[lab]+[fmt(abl[k][i]) for i in (0,1,2,3)] for k,lab in amap])
para('El número de escalas (n) es el factor determinante; los elementos estructurantes lineales aportan poco frente al disco solo en estas métricas.', size=10, after=10)

heading('1.6 Evaluación orientada a tarea — detección con YOLOv8n (inferencia)')
disp={k:n for k,n in ORDER}; disp['VIS']='VIS (solo visible)'; disp['IR']='IR (solo infrarrojo)'
rows=[]
for k in ['VIS','IR']+[o[0] for o in ORDER]:
    if k in ds.index:
        rows.append([disp.get(k,k),int(ds.loc[k,'det_total']),int(ds.loc[k,'person_total']),int(ds.loc[k,'vehiculo_total']),fmt(ds.loc[k,'conf_media']),fmt(ds.loc[k,'FPS'],1)])
table(['Método','Detecciones','Personas','Vehículos','Conf. media','FPS'],rows)
para('La fusión mejora la detectabilidad de personas frente a VIS. Las diferencias entre métodos no son significativas (n=20 sin etiquetas); '
 'un mAP concluyente requiere un conjunto etiquetado.', size=10, after=8)

# ---- cualitativo ----
doc.add_page_break()
heading('2. Comparación cualitativa por escena', 14)
para('Para cada una de las 20 escenas se muestran las imágenes fuente (VIS, IR) y la salida de los 13 métodos de fusión, '
 'para observar el efecto visual de cada uno.', after=8)
monts=sorted(glob.glob('docs/figures/cualitativas/montaje_*.png'))
for i,mp in enumerate(monts,1):
    doc.add_picture(mp, width=Inches(6.5)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if i%2==0 and i<len(monts): doc.add_page_break()

doc.save('/tmp/resultados_plain.docx')
print('OK', len(doc.tables),'tablas,',len(monts),'montajes')

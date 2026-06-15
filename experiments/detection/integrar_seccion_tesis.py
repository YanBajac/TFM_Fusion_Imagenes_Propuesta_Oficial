# -*- coding: utf-8 -*-
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT=Path(".")
DOC=ROOT/"docs"/"Tesis_Borrador.docx"
FIG=ROOT/"docs"/"figures"/"fig_deteccion_yolo.png"
s=pd.read_csv(ROOT/"experiments"/"results"/"metrics_reports"/"detection_summary.csv").set_index("method")

def grp(m):
    if m in("VIS","IR"): return m
    if m in("Promedio","PiramideLaplace","Curvelet"): return "Baselines"
    return "Fusión +BTH" if "BTH" in m else "Fusión WTH"
s["g"]=[grp(m) for m in s.index]
order=["VIS","IR","Baselines","Fusión WTH","Fusión +BTH"]
agg=s.groupby("g").agg(det=("det_total","sum"),per=("person_total","sum"),
     veh=("vehiculo_total","sum"),conf=("conf_media","mean"),fps=("FPS","mean")).reindex(order).round(2)

d=Document(str(DOC))
H2=[p.style for p in d.paragraphs if p.style and p.style.name=="Heading 2"][0]
def find(prefix):
    for p in d.paragraphs:
        if p.text.strip().startswith(prefix): return p
def ins_after(ref,text="",style=None):
    e=OxmlElement("w:p"); ref._p.addnext(e); np_=Paragraph(e,ref._parent)
    if style is not None: np_.style=style
    if text: np_.add_run(text)
    return np_
def set_cell(c,t,b=False):
    for r in list(c.paragraphs[0].runs): r._r.getparent().remove(r._r)
    run=c.paragraphs[0].add_run(t); run.bold=b

# anclar después de la Figura 6 (BTH); si no, después de discusión
anchor=find("Figura 6.") or find("Los hallazgos cuantitativos pueden sintetizarse")
cur=ins_after(anchor,"Evaluación orientada a tarea: detección con YOLO",style=H2)
p1=ins_after(cur,
 "Para evaluar la utilidad de la fusión más allá de la calidad de imagen se realizó una evaluación "
 "orientada a tarea con un detector YOLOv8n preentrenado (clases COCO). Sin reentrenar ni anotar, se "
 "ejecutó inferencia sobre las 20 imágenes en cada versión (VIS, IR y las once fusiones) y se contabilizó "
 "la detectabilidad: número de detecciones, clase (persona/vehículo) y confianza. La Tabla 6 agrupa los "
 "resultados por tipo de fuente.")
# tabla
cap=ins_after(p1,"Tabla 6. Detectabilidad de YOLOv8n por tipo de fuente (suma sobre 20 imágenes).")
t=d.add_table(rows=1,cols=6)
try: t.style=d.tables[3].style
except Exception: pass
cap._p.addnext(t._tbl)
hdr=["Fuente","Detecciones","Personas","Vehículos","Confianza media","FPS (CPU)"]
for j,h in enumerate(hdr): set_cell(t.rows[0].cells[j],h,b=True)
for g in order:
    row=t.add_row().cells
    set_cell(row[0],g)
    set_cell(row[1],str(int(agg.loc[g,"det"])))
    set_cell(row[2],str(int(agg.loc[g,"per"])))
    set_cell(row[3],str(int(agg.loc[g,"veh"])))
    set_cell(row[4],f"{agg.loc[g,'conf']:.2f}")
    set_cell(row[5],f"{agg.loc[g,'fps']:.1f}")

# texto de interpretación (después de la tabla)
p2=ins_after(cap,
 "El hallazgo principal respalda la utilidad de la fusión: sobre la VIS sola el detector apenas encuentra "
 "una (1) persona en todo el conjunto —coherente con la baja visibilidad de las personas en el espectro "
 "visible bajo humo o de noche—, mientras que al incorporar la información infrarroja mediante la fusión "
 "el número de personas detectadas asciende a trece (Top-Hat WTH) y dieciséis (variantes con BTH). "
 "Emerge además un patrón complementario coherente con la física de cada sensor: la VIS favorece la "
 "detección de vehículos (textura visible) y la IR la de personas (firma térmica). La fusión Top-Hat WTH "
 "es la única que equilibra ambas clases (13 personas y 12 vehículos), mientras que las variantes con "
 "Black Top-Hat elevan la detección de personas pero degradan la de vehículos (4), de forma consistente "
 "con los artefactos que el BTH introduce en los bordes de objetos extensos.")
# la figura
img=ins_after(p2,"")
img.add_run().add_picture(str(FIG), width=Inches(6.0))
fcap=ins_after(img,"Figura 7. Detecciones de YOLOv8n por modalidad (izq.: personas vs vehículos por grupo; der.: detecciones totales por método).")
# limitaciones
ins_after(fcap,
 "Esta evaluación debe leerse con cautela. Se trata de detectabilidad sin verdad de campo (ground truth): "
 "un mayor número de detecciones no equivale necesariamente a mayor exactitud, pues podría incluir falsos "
 "positivos. Con solo 20 imágenes sin etiquetar, la prueba de Friedman sobre las métricas de detección no "
 "resultó significativa (p = 0,29 para el número de detecciones; p = 0,79 para la confianza), por lo que no "
 "es posible declarar un método de fusión ganador en términos de detección. La inferencia en CPU alcanzó "
 "1,6–2,7 FPS. Para una comparación concluyente en términos de mAP, precisión, recall y F1 se requiere un "
 "conjunto etiquetado (p. ej. mediante Roboflow o el dataset M3FD) y un mayor número de muestras; el "
 "procedimiento y los notebooks correspondientes se entregan como base para ese trabajo (Apéndice y "
 "repositorio).")

# siglas nuevas
t0=d.tables[0]
existing={r.cells[0].text.strip() for r in t0.rows}
for sig,txt in [("YOLO","You Only Look Once (detector de objetos)"),
                ("FPS","Cuadros por segundo (Frames Per Second)"),
                ("mAP","Precisión media promedio (mean Average Precision)")]:
    if sig not in existing:
        c=t0.add_row().cells; set_cell(c[0],sig); set_cell(c[1],txt)

d.save(str(DOC))
print("Sección de detección integrada. Tabla por grupo:")
print(agg.to_string())

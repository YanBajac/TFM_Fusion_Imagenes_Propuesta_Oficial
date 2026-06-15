# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
d=Document('docs/Tesis_Borrador.docx')
P=d.paragraphs

def set_text(par,text):
    for r in list(par.runs): r._r.getparent().remove(r._r)
    par.add_run(text)
def find(prefix):
    for p in P:
        if p.text.strip().startswith(prefix): return p
    raise KeyError(prefix)
def insert_after(ref,text,style=None):
    np_=OxmlElement('w:p'); ref._p.addnext(np_)
    P2=Paragraph(np_,ref._parent)
    if style is not None: P2.style=style
    if text: P2.add_run(text)
    return P2

# 1) SUMMARY en ingles (espejo del resumen actualizado)
sm=find("SUMMARY")
# el parrafo de contenido es el siguiente no vacio
body=Paragraph(sm._p.getnext(),sm._parent)
set_text(body,
 "Image fusion in the visible (VIS) and infrared (IR) spectra is a key technique for computer "
 "perception under adverse conditions, with applications in surveillance, autonomous driving and "
 "industrial inspection. This work proposes and evaluates the Top-Hat Tower: a multi-scale "
 "morphological decomposition with configurable structuring elements (disk, square, cross), variable "
 "number of levels and base radii, in two modes: White Top-Hat (WTH) and White+Black Top-Hat (WTH+BTH). "
 "Detail layers are fused via local-activity selection while the low-frequency base is averaged. "
 "Experiments were carried out over twenty (20) VIS/IR pairs from the TNO Image Fusion Dataset, "
 "comparing the Top-Hat Tower against three baselines (average, Laplacian pyramid and curvelet). Quality "
 "was measured with twelve no-reference metrics: six classical metrics of activity and information "
 "(EN, SD, FE, MG, MI_vis, MI_ir) and six standard fusion-quality metrics (SF, Qabf, Nabf, SSIM, SCD, "
 "VIF). Non-parametric analysis (Friedman, paired Wilcoxon with Holm correction and effect size, mean "
 "rank) revealed significant differences (p < 0.001 in all twelve metrics). The Laplacian pyramid leads "
 "the aggregate ranking (4.42) thanks to its advantage in contrast and perceptual fidelity, but the best "
 "Top-Hat WTH configuration (disk, L = 5; rank 5.00) significantly outperforms it in structural fidelity "
 "(SSIM, SCD) and matches it in edge preservation (Qabf) and artifacts (Nabf): no single method "
 "dominates, each being optimal under complementary criteria. The Black Top-Hat variant increases "
 "contrast (EN, SD, SF, MG) but degrades quality (Qabf, SSIM, VIF) and multiplies artifacts (Nabf), and "
 "is therefore not recommended by default. The Top-Hat Tower is concluded to be a classical, "
 "interpretable, computationally light method, superior to the Laplacian pyramid in structural fidelity "
 "to the sources.")

# 2) SIGLAS: agregar 6 metricas
t0=d.tables[0]
for sig,sig_txt in [
    ("SF","Frecuencia espacial (Spatial Frequency)"),
    ("Qabf","Métrica de preservación de bordes (Xydeas-Petrovic)"),
    ("Nabf","Artefactos añadidos por la fusión (menor es mejor)"),
    ("SSIM","Índice de similitud estructural (Structural Similarity)"),
    ("SCD","Suma de correlaciones de las diferencias"),
    ("VIF","Fidelidad de información visual (Visual Information Fidelity)"),
]:
    cells=t0.add_row().cells
    set_text(cells[0].paragraphs[0],sig)
    set_text(cells[1].paragraphs[0],sig_txt)

# 3) REFERENCIAS: agregar 4 (en orden alfabetico aproximado)
ref_style=find("Bai, X. (2013)").style
insert_after(find("REFERENCIAS BIBLIOGRÁFICAS"),
    "Aslantas, V., y Bendes, E. (2015). A new image quality metric for image fusion: The sum of the "
    "correlations of differences. AEU - International Journal of Electronics and Communications, 69(12), "
    "1890–1896.", style=ref_style)
insert_after(find("Gonzalez, R. C."),
    "Kumar, B. K. S. (2013). Multifocus and multispectral image fusion based on pixel significance using "
    "discrete cosine harmonic wavelet transform. Signal, Image and Video Processing, 7(6), 1125–1143.",
    style=ref_style)
insert_after(find("Serra, J. (1982)"),
    "Sheikh, H. R., y Bovik, A. C. (2006). Image information and visual quality. IEEE Transactions on "
    "Image Processing, 15(2), 430–444.", style=ref_style)
insert_after(find("Wilcoxon, F. (1945)"),
    "Xydeas, C. S., y Petrović, V. (2000). Objective image fusion performance measure. Electronics "
    "Letters, 36(4), 308–309.", style=ref_style)

# 4) APENDICE C: actualizar conteo de contrastes
set_text(find("Las tablas completas con los"),
    "Las tablas completas con los 288 contrastes Wilcoxon (8 configuraciones Top-Hat × 3 baselines × 12 "
    "métricas) se almacenan en experiments/results/metrics_reports/wilcoxon_results.csv. Cada fila "
    "contiene la métrica, las medias comparadas, la diferencia, el estadístico W, el p-valor, el p-valor "
    "corregido por Holm, la marca de significación a α = 0,05 y el tamaño de efecto rank-biserial. El "
    "ranking promedio está en ranking_methods.csv y los resultados de Friedman en friedman_results.csv.")

d.save('docs/Tesis_Borrador.docx')
print("Gaps del libro corregidos: SUMMARY, siglas, 4 referencias, apendice C.")

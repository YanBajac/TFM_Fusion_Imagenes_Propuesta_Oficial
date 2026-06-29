# -*- coding: utf-8 -*-
import copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

d=Document('/tmp/base.docx')
H1=next(p.style for p in d.paragraphs if p.style and p.style.name=='Heading 1')

# ---------- 1) Renumerar tablas 1..10 (token remap) ----------
MAP=[('Tabla 3a','@@3@@'),('Tabla 3b','@@5@@'),('Tabla 3','@@4@@'),
     ('Tabla 4','@@6@@'),('Tabla 5','@@7@@'),('Tabla 6','@@8@@'),
     ('Tabla 7','@@9@@'),('Tabla 8','@@10@@')]
def remap(txt):
    for a,b in MAP: txt=txt.replace(a,b)
    for n in [10,9,8,7,6,5,4,3]: txt=txt.replace(f'@@{n}@@',f'Tabla {n}')
    return txt
for p in d.paragraphs:
    if 'Tabla ' in p.text:
        new=remap(p.text)
        if new!=p.text:
            cap=p.text.strip().startswith('Tabla ')
            for r in list(p.runs): r._r.getparent().remove(r._r)
            run=p.add_run(new)
            if cap: run.italic=True; run.font.size=Pt(9)

# ---------- 2) Captions faltantes (clásicas=4, Friedman=6, ranking=7) ----------
def cap_before(tbl, text):
    pe=OxmlElement('w:p'); tbl._tbl.addprevious(pe)
    pp=Paragraph(pe,tbl._parent); r=pp.add_run(text); r.italic=True; r.font.size=Pt(9)
def header_of(t): return ' | '.join(c.text[:8] for c in t.rows[0].cells)
for t in d.tables:
    h=header_of(t)
    if h.startswith('Método | EN'):   cap_before(t,'Tabla 4. Promedios de las seis métricas clásicas por método (n = 20).')
    elif h.startswith('Métrica | χ²') or h.startswith('Métrica | χ'): cap_before(t,'Tabla 6. Prueba de Friedman por métrica (χ², p-valor; α = 0,05).')
    elif h.startswith('Método | SD | MG'): cap_before(t,'Tabla 7. Ranking promedio por método y métrica (1 = mejor).')

# ---------- 3) LISTA DE FIGURAS y LISTA DE TABLAS antes de INTRODUCCIÓN ----------
intro=next(p for p in d.paragraphs if p.text.strip()=='INTRODUCCIÓN')
figs=[p.text.strip() for p in d.paragraphs if p.text.strip().startswith('Figura ')]
tabs=[p.text.strip() for p in d.paragraphs if p.text.strip().startswith('Tabla ')]
import re
def keynum(s):
    m=re.match(r'(Figura|Tabla)\s+(\d+)',s); return int(m.group(2)) if m else 999
figs=sorted(set(figs),key=keynum); tabs=sorted(set(tabs),key=keynum)

def add_heading_before(anchor,text):
    pe=OxmlElement('w:p'); anchor._p.addprevious(pe)
    pp=Paragraph(pe,anchor._parent); pp.style=H1; pp.add_run(text); return pp
def add_line_before(anchor,text):
    pe=OxmlElement('w:p'); anchor._p.addprevious(pe)
    pp=Paragraph(pe,anchor._parent); r=pp.add_run(text); r.font.size=Pt(10.5)
    pp.paragraph_format.space_after=Pt(2); return pp

add_heading_before(intro,'LISTA DE FIGURAS')
for f in figs: add_line_before(intro,f)
add_heading_before(intro,'LISTA DE TABLAS')
for t in tabs: add_line_before(intro,t)

# ---------- 4) Pseudocódigo multiescala en Apéndice D ----------
# encontrar heading "E." (siguiente apéndice tras D) para insertar antes
apnd_E=None
for p in d.paragraphs:
    if p.style and p.style.name=='Heading 2' and re.match(r'^E\.\s', p.text.strip()):
        apnd_E=p; break
PSEUDO=[
 'ENTRADA: I_VIS, I_IR en [0,1] alineadas;  n escalas;  radio base r;  peso m',
 'SALIDA:  imagen fusionada F',
 '',
 'para cada fuente S en {I_VIS, I_IR}:',
 '    para i = 1 ... n:                      # escalas crecientes',
 '        r_i  <- redondear(r * i)',
 '        SE   <- {disco(r_i), linea(r_i,0), linea(r_i,45), linea(r_i,90), linea(r_i,135)}',
 '        WTH_i <- max_{b en SE} ( S - apertura(S,b) )        # ec. (4)',
 '        BTH_i <- max_{b en SE} ( cierre(S,b) - S )',
 '    WTHS_2 <- WTH_2 - WTH_1 ;  BTHS_2 <- BTH_2 - BTH_1       # ec. (5)',
 '    para i = 3 ... n:',
 '        WTHS_i <- WTH_i - WTHS_{i-2} ;  BTHS_i <- BTH_i - BTHS_{i-2}',
 '    WTH_M[S]  <- max_i WTH_i ;  WTHS_M[S] <- max_i WTHS_i    # ec. (6)',
 '    BTH_M[S]  <- max_i BTH_i ;  BTHS_M[S] <- max_i BTHS_i',
 '',
 '# combinacion entre fuentes (maximo por pixel)',
 'WTH_M  <- max(WTH_M[VIS],  WTH_M[IR]) ;   WTHS_M <- max(WTHS_M[VIS], WTHS_M[IR])',
 'BTH_M  <- max(BTH_M[VIS],  BTH_M[IR]) ;   BTHS_M <- max(BTHS_M[VIS], BTHS_M[IR])',
 'I_base <- (I_VIS + I_IR) / 2',
 'F <- recortar( I_base + m*(WTH_M + WTHS_M) - m*(BTH_M + BTHS_M), 0, 1 )   # ec. (7)',
 'devolver F',
 '',
 '(n, r, m) se optimizan por PSO  ->  n = 6, r ~ 2,89, m = 0,10',
]
def ins_block_before(anchor):
    # subtitulo
    pe=OxmlElement('w:p'); anchor._p.addprevious(pe)
    pp=Paragraph(pe,anchor._parent); r=pp.add_run('Pseudocódigo del método óptimo multiescala'); r.bold=True
    last=pp
    for line in PSEUDO:
        e=OxmlElement('w:p'); anchor._p.addprevious(e)
        q=Paragraph(e,anchor._parent); rr=q.add_run(line if line else ' ')
        rr.font.name='Courier New'; rr.font.size=Pt(8.5)
        q.paragraph_format.space_after=Pt(0); q.paragraph_format.space_before=Pt(0)
if apnd_E is not None:
    ins_block_before(apnd_E)
else:
    print('WARN: no se encontró apéndice E; pseudocódigo no insertado')

d.save('/tmp/out30.docx')
import zipfile; print('out30 zip ok:',zipfile.is_zipfile('/tmp/out30.docx'),'| figs',len(figs),'tabs',len(tabs))

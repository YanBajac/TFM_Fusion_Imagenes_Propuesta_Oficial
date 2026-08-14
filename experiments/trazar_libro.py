# -*- coding: utf-8 -*-
"""Rastrea cada cifra del libro hasta el CSV que la produce, y lista las que no salen de ninguno.

PARA QUE SIRVE. El libro se revisa a mano para comprobar que ninguna cifra fue inventada. Verificar
las seiscientas cifras de setenta y cinco paginas a mano es inviable; verificar las que este script no
puede rastrear, si. La salida separa tres cosas:

  TRAZADA ....... el valor aparece en al menos un CSV de experiments/results/, con el mismo redondeo.
                  No prueba que la FRASE sea correcta —para eso hay que leerla— pero si que el numero
                  existe en los datos y no salio de la nada.
  DECLARADA ..... es un valor de diseño o de convencion, no un resultado: el radio adoptado, el nivel
                  de significancia, un año de publicacion, el tamaño de un elemento estructurante.
                  Van con su motivo, uno por uno, para que la lista no sea una bolsa de excepciones.
  SIN FUENTE .... no esta en ningun CSV ni en la lista de declaradas. ES LA COLUMNA QUE HAY QUE MIRAR.

LO QUE NO HACE, y conviene tenerlo claro. No verifica que la afirmacion que rodea a la cifra sea
correcta: «la propuesta queda 1.a con 3,394» y «la propuesta queda 7.a con 3,394» son igual de
trazables. Tampoco verifica los conteos verbales —«supera a seis de las siete»— porque no son cifras.
Eso se lee.

Uso:  .venv\\Scripts\\python.exe -X utf8 experiments/trazar_libro.py [--todas]
      --todas  imprime tambien las trazadas, no solo las que hay que revisar
"""
import os
import re
import sys
from collections import defaultdict
from pathlib import Path

RAIZ = Path(__file__).resolve().parent.parent
os.chdir(RAIZ)
import pandas as pd
from docx import Document

LIBRO = RAIZ / 'docs' / 'Tesis_Borrador_V3.docx'
RES = RAIZ / 'experiments' / 'results'
TODAS = '--todas' in sys.argv

# Valores que no son un resultado sino una decision, una convencion o un dato bibliografico. Cada uno
# va con su motivo: sin el, esta lista se convierte en el lugar donde se esconde lo que no cuadra.
DECLARADOS = {
    '25': 'el radio adoptado, r = 25 (y el tamaño de la rejilla del PSO)',
    '0,30': 'el peso adoptado, m = 0,30',
    '51': 'el diametro del elemento estructurante con r = 25: 2·25 + 1',
    '0,05': 'el nivel de significancia α = 0,05',
    '2,00': 'el techo del rango de busqueda del peso, heredado de la referencia',
    '0,5': 'el umbral de IoU del mAP@0,5',
    '0,95': 'el techo del mAP@0,5:0,95',
    '3': 'el tamaño 3×3 del elemento minimo, y el numero de terminos de la aptitud',
    '5': 'los cinco elementos estructurantes del banco, y las cinco semillas',
    '7': 'los siete metodos del benchmark',
    '9': 'las nueve metricas reportadas',
    '17': 'las diecisiete metricas que el evaluador computa',
    '20': 'los veinte pares del TNO',
    '40': 'las epocas de entrenamiento del detector',
    '2.000': 'las imagenes de entrenamiento de LLVIP',
    '500': 'las imagenes de validacion de LLVIP, y las 500 corridas del estudio del PSO',
    '2025': 'año de publicacion (Ortega y Espinoza)',
    '2024': 'año de publicacion',
    '2023': 'año de publicacion',
    '2022': 'año de publicacion',
    '2021': 'año de publicacion',
    '2020': 'año de publicacion',
    '2019': 'año de publicacion',
    '2017': 'año de publicacion (Toet, dataset TNO)',
    '2016': 'año de publicacion',
    '2026': 'año de la defensa',
    # Propiedades de los datasets y de las imagenes: no las produce este trabajo, y por eso no hay CSV
    # donde buscarlas. Se verifican contra la fuente del dataset, no contra un resultado propio.
    '15.488': 'el tamaño completo de LLVIP segun su publicacion: 15.488 pares',
    '576': 'la resolucion mayor del corpus TNO: 768×576 pixeles',
    '270': 'la resolucion menor del corpus TNO: 360×270 pixeles',
    '360': 'la resolucion menor del corpus TNO: 360×270 pixeles',
    '768': 'la resolucion mayor del corpus TNO: 768×576 pixeles',
    '640': 'el lado de entrada del detector (imgsz = 640)',
    # Derivada de un parametro, no medida: prepare_m3fd.py usa --train-n 2000 y el detector unico se
    # entrena con las dos modalidades mezcladas, o sea 2 × 2.000.
    '4.000': 'las imagenes de entrenamiento de M3FD: 2.000 pares por las dos modalidades',
}


def numeros(texto):
    """Las cifras de un texto, tal como estan escritas, para no perder el redondeo.

    El separador decimal se decide POR PARRAFO y no para todo el libro: el resumen esta en español y
    usa coma, y el abstract esta en ingles y usa punto, asi que «0.906» en el abstract es un mAP y
    «2.000» en el cuerpo son dos mil imagenes. Mirando el parrafo se distingue: si tiene coma decimal
    es español, y si tiene punto decimal y ninguna coma decimal, es ingles. Sin esto, las nueve cifras
    del abstract salian como cifras sin fuente y tapaban las que si hay que revisar.
    """
    ingles = re.search(r'\d\.\d', texto) and not re.search(r'\d,\d', texto)
    if ingles:
        toks = re.findall(r'\d+\.\d+|\d+', texto)
        return [t.replace('.', ',') for t in toks]
    return re.findall(r'\d{1,3}(?:\.\d{3})+,\d+|\d+,\d+|\d{1,3}(?:\.\d{3})+|\d+', texto)


def a_float(tok):
    t = tok.replace('.', '').replace(',', '.') if ',' in tok else tok.replace('.', '')
    try:
        return float(t)
    except ValueError:
        return None


def cargar_csv():
    """Todos los valores numericos de todos los CSV, indexados por su forma redondeada.

    Se indexa por CADENA y no por float a proposito: el libro escribe «3,394» y el CSV guarda
    3.3941176470588235, asi que lo que hay que comparar es el valor del CSV redondeado a los decimales
    que el libro usa. Se guardan los redondeos de 0 a 4 decimales, mas la version en porcentaje, que es
    la otra forma en que un mismo dato aparece escrito (0,50 en el CSV y 50,0 % en el texto).
    """
    idx = defaultdict(set)
    n_arch = 0
    for p in sorted(RES.rglob('*.csv')):
        try:
            df = pd.read_csv(p)
        except Exception:
            continue
        n_arch += 1
        rel = str(p.relative_to(RES)).replace('\\', '/')
        vals = []
        for col in df.columns:
            s = pd.to_numeric(df[col], errors='coerce').dropna()
            vals.extend(s.tolist())
        # los nombres de columna tambien llevan cifras (mAP50, 3x3) pero no son datos: no se indexan
        for v in vals:
            for nd in range(5):
                idx[f'{v:.{nd}f}'.replace('.', ',')].add(rel)
                idx[f'{v * 100:.{nd}f}'.replace('.', ',')].add(rel)
    return idx, n_arch


def main():
    idx, n_arch = cargar_csv()
    print(f'=== {n_arch} CSV indexados desde experiments/results/ ===\n')

    doc = Document(str(LIBRO))

    # La bibliografia se saltea entera, y hay que decir por que: sus DOI («10.1016/j.inffus...») y sus
    # numeros de volumen y pagina no son resultados, y al trocearlos el rastreador escupia cuarenta
    # cifras sin fuente que no hay nada que revisar. Los limites se BUSCAN por encabezado y no se
    # escriben a mano, para que sigan valiendo si el capitulo cambia de lugar.
    def _titulos(patron):
        """Los ENCABEZADOS que matchean, no cualquier parrafo que los mencione.

        Sin el filtro de estilo esto agarraba la entrada del indice —«7. REFERENCIAS BIBLIOGRAFICAS»
        aparece tambien en la lista de contenidos— y tomaba el parrafo 160 como comienzo del capitulo
        en vez del 443, de modo que la bibliografia no se salteaba y sus DOI seguian saliendo como
        cifras sin fuente.
        """
        return [i for i, p in enumerate(doc.paragraphs)
                if re.match(patron, p.text.strip(), re.I)
                and (p.style.name or '').startswith(('Heading 1', 'Título 1', 'Titulo 1'))]

    _ini_bib = (_titulos(r'\d+\.\s*REFERENCIAS') or [None])[0]
    _fin_bib = (_titulos(r'\d+\.\s*AP[ÉE]NDICE') or [None])[0]
    salteados = set()
    if _ini_bib is not None and _fin_bib is not None and _ini_bib < _fin_bib:
        salteados = set(range(_ini_bib, _fin_bib))
        print(f'  se saltea la bibliografia: parrafos {_ini_bib} a {_fin_bib - 1} '
              f'(DOI, volumenes y paginas, que no son resultados)\n')

    # de donde sale cada cifra del libro: parrafo N, o tabla N
    apariciones = defaultdict(list)
    for i, p in enumerate(doc.paragraphs):
        if i in salteados:
            continue
        for tok in numeros(p.text):
            apariciones[tok].append(f'par. {i}')
    for ti, t in enumerate(doc.tables):
        for fila in t.rows:
            for c in fila.cells:
                for tok in numeros(c.text):
                    apariciones[tok].append(f'tabla {ti}')

    trazadas, declaradas, huerfanas = {}, {}, {}
    for tok, donde in apariciones.items():
        v = a_float(tok)
        if tok in DECLARADOS:
            declaradas[tok] = (donde, DECLARADOS[tok])
        elif v is not None and ',' not in tok and 1900 <= v <= 2030 and v == int(v):
            # un entero de cuatro cifras en ese rango es un año de publicacion, no un resultado. Se
            # resuelve por regla y no agregando veinte años a la lista de declarados uno por uno.
            declaradas[tok] = (donde, 'año de publicacion o de la defensa')
        elif tok in idx:
            trazadas[tok] = (donde, sorted(idx[tok]))
        else:
            huerfanas[tok] = donde

    n_ap = sum(len(v) for v in apariciones.values())
    print(f'  {len(apariciones)} cifras distintas, {n_ap} apariciones en el libro')
    print(f'    TRAZADAS   {len(trazadas):>4} valores distintos, en algun CSV')
    print(f'    DECLARADAS {len(declaradas):>4} valores de diseño o bibliograficos')
    print(f'    SIN FUENTE {len(huerfanas):>4} valores que NO salen de ningun CSV')

    print('\n' + '=' * 96)
    print('SIN FUENTE — las cifras que hay que revisar a mano')
    print('=' * 96)
    # ordenadas por cuantas veces aparecen: la que se repite es la que mas daño hace si esta mal
    for tok, donde in sorted(huerfanas.items(), key=lambda kv: (-len(kv[1]), kv[0])):
        v = a_float(tok)
        pista = ''
        if v is not None and v == int(v) and 1 <= v <= 999:
            pista = '  (entero chico: probablemente un conteo, una remision o un numero de ecuacion)'
        print(f'  {tok:>12}  x{len(donde):<3} {", ".join(sorted(set(donde))[:6])}{pista}')

    print('\n' + '=' * 96)
    print('DECLARADAS — no son resultados, y por eso no estan en ningun CSV')
    print('=' * 96)
    for tok, (donde, motivo) in sorted(declaradas.items(), key=lambda kv: -len(kv[1][0])):
        print(f'  {tok:>12}  x{len(donde):<3} {motivo}')

    if TODAS:
        print('\n' + '=' * 96)
        print('TRAZADAS — valor, donde aparece, y en que CSV esta')
        print('=' * 96)
        for tok, (donde, arch) in sorted(trazadas.items(), key=lambda kv: (-len(kv[1][0]), kv[0])):
            print(f'  {tok:>12}  x{len(donde):<3} {", ".join(arch[:3])}'
                  + (f' (+{len(arch) - 3})' if len(arch) > 3 else ''))
    else:
        print(f'\n(las {len(trazadas)} trazadas se listan con --todas)')


if __name__ == '__main__':
    main()
